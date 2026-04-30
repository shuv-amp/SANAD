import csv
import io
import json
import zipfile
from pathlib import Path

import fitz
import pytest
from docx import Document
from fastapi.testclient import TestClient
from sqlalchemy import create_engine
from sqlalchemy.orm import Session, sessionmaker
from sqlalchemy.pool import StaticPool

from sanad_api.config import Settings, get_settings
from sanad_api.database import Base, get_db
from sanad_api.main import create_app
from sanad_api.services.glossary import seed_default_glossary

SAMPLES_DIR = Path(__file__).resolve().parents[3] / "samples" / "demo"
ROOT_DIR = Path(__file__).resolve().parents[3]


@pytest.fixture
def client(tmp_path: Path, monkeypatch: pytest.MonkeyPatch):
    monkeypatch.setenv("SANAD_STORAGE_ROOT", str(tmp_path / "storage"))
    monkeypatch.setenv("SANAD_ACTIVE_PROVIDER", "fixture")
    get_settings.cache_clear()

    engine = create_engine(
        "sqlite://",
        connect_args={"check_same_thread": False},
        poolclass=StaticPool,
        future=True,
    )
    Base.metadata.create_all(engine)
    session_local = sessionmaker(bind=engine, autoflush=False, autocommit=False, future=True)
    with Session(engine) as db:
        seed_default_glossary(db)

    def override_db():
        db = session_local()
        try:
            yield db
        finally:
            db.close()

    app = create_app()
    app.dependency_overrides[get_db] = override_db
    app.state.session_local = session_local
    with TestClient(app) as test_client:
        yield test_client
    get_settings.cache_clear()


def test_invalid_pdf_fails_during_processing_with_explicit_message(client: TestClient) -> None:
    upload = client.post(
        "/api/documents",
        data={"source_lang": "en", "target_lang": "ne", "domain": "public_service", "subdomain": "residence"},
        files={"file": ("scan.pdf", b"%PDF scanned placeholder", "application/pdf")},
    )

    assert upload.status_code == 200
    document_id = upload.json()["id"]

    response = client.post(f"/api/documents/{document_id}/process")

    assert response.status_code == 400
    assert "Could not parse this PDF" in response.json()["detail"]


def test_debug_reset_demo_clears_current_state(client: TestClient, monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setattr("sanad_api.services.demo_reset._regenerate_demo_fixtures", lambda: None)

    fixture = SAMPLES_DIR / "public-service-1.docx"
    with fixture.open("rb") as handle:
        upload = client.post(
            "/api/documents",
            data={"source_lang": "en", "target_lang": "ne", "domain": "public_service", "subdomain": "residence"},
            files={
                "file": (
                    fixture.name,
                    handle,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )

    document_id = upload.json()["id"]
    process = client.post(f"/api/documents/{document_id}/process")
    assert process.status_code == 200

    reset = client.post("/api/debug/reset-demo")
    assert reset.status_code == 200
    assert reset.json()["status"] == "reset"
    assert reset.json()["fixtures_regenerated"] is True

    missing = client.get(f"/api/documents/{document_id}")
    assert missing.status_code == 404


def test_plain_text_document_processes_and_exports_as_docx(client: TestClient, tmp_path: Path) -> None:
    source = tmp_path / "notice.txt"
    source.write_text(
        "Certificate of Residence Request\n\n"
        "Please submit this form to the Ward Office.\n\n"
        "Fee: NPR 500\n",
        encoding="utf-8",
    )

    with source.open("rb") as handle:
        upload = client.post(
            "/api/documents",
            data={"source_lang": "en", "target_lang": "ne", "domain": "public_service", "subdomain": "residence"},
            files={"file": ("notice.txt", handle, "text/plain")},
        )

    assert upload.status_code == 200
    document_id = upload.json()["id"]

    process = client.post(f"/api/documents/{document_id}/process")
    assert process.status_code == 200
    assert process.json()["file_type"] == "txt"

    segments = client.get(f"/api/documents/{document_id}/segments").json()
    assert [segment["source_text"] for segment in segments] == [
        "Certificate of Residence Request",
        "Please submit this form to the Ward Office.",
        "Fee: NPR 500",
    ]

    for segment in segments:
        response = client.post(
            f"/api/segments/{segment['id']}/approve",
            json={"text": segment["translation"]["candidate_text"], "actor": "test-reviewer"},
        )
        assert response.status_code == 200

    export = client.post(f"/api/documents/{document_id}/export", json={"format": "docx"})
    assert export.status_code == 200

    exported = Document(export.json()["export_file_uri"])
    exported_lines = [paragraph.text for paragraph in exported.paragraphs if paragraph.text.strip()]
    assert "बसोबास प्रमाणपत्र अनुरोध" in exported_lines
    assert "कृपया यो फारम वडा कार्यालयमा बुझाउनुहोस्।" in exported_lines
    assert "शुल्क: NPR ५००" in exported_lines


def test_csv_document_processes_and_exports_in_same_format(client: TestClient, tmp_path: Path) -> None:
    source = tmp_path / "notice.csv"
    source.write_text(
        "Field,Value\n"
        "Fee,NPR 500\n"
        "Date,2026-04-21\n",
        encoding="utf-8",
    )

    with source.open("rb") as handle:
        upload = client.post(
            "/api/documents",
            data={"source_lang": "en", "target_lang": "ne", "domain": "public_service", "subdomain": "residence"},
            files={"file": ("notice.csv", handle, "text/csv")},
        )

    assert upload.status_code == 200
    document_id = upload.json()["id"]

    process = client.post(f"/api/documents/{document_id}/process")
    assert process.status_code == 200
    assert process.json()["file_type"] == "csv"

    segments = client.get(f"/api/documents/{document_id}/segments").json()
    assert [segment["source_text"] for segment in segments] == [
        "Field",
        "Value",
        "Fee",
        "NPR 500",
        "Date",
        "2026-04-21",
    ]

    for segment in segments:
        response = client.post(
            f"/api/segments/{segment['id']}/approve",
            json={"text": segment["translation"]["candidate_text"], "actor": "test-reviewer"},
        )
        assert response.status_code == 200

    export = client.post(f"/api/documents/{document_id}/export", json={"format": "csv"})
    assert export.status_code == 200
    assert export.json()["format"] == "csv"

    exported_text = Path(export.json()["export_file_uri"]).read_text(encoding="utf-8").splitlines()
    assert exported_text[0]
    assert len(exported_text) == 3
    assert "NPR ५००" in exported_text[1]


def test_pdf_document_processes_and_exports_in_same_format(client: TestClient, tmp_path: Path) -> None:
    source = tmp_path / "notice.pdf"
    _create_pdf(
        source,
        [
            "Residence Certificate Notice",
            "Date: 2026-05-02",
            "Fee: NPR 500",
            "Reference ID: RES-2026-004",
        ],
    )

    with source.open("rb") as handle:
        upload = client.post(
            "/api/documents",
            data={"source_lang": "en", "target_lang": "ne", "domain": "public_service", "subdomain": "residence"},
            files={"file": ("notice.pdf", handle, "application/pdf")},
        )

    assert upload.status_code == 200
    document_id = upload.json()["id"]

    process = client.post(f"/api/documents/{document_id}/process")
    assert process.status_code == 200
    assert process.json()["file_type"] == "pdf"

    segments = client.get(f"/api/documents/{document_id}/segments").json()
    assert [segment["source_text"] for segment in segments] == [
        "Residence Certificate Notice",
        "Date: 2026-05-02",
        "Fee: NPR 500",
        "Reference ID: RES-2026-004",
    ]

    for segment in segments:
        response = client.post(
            f"/api/segments/{segment['id']}/approve",
            json={"text": segment["translation"]["candidate_text"], "actor": "test-reviewer"},
        )
        assert response.status_code == 200

    export = client.post(f"/api/documents/{document_id}/export", json={"format": "pdf"})
    assert export.status_code == 200
    assert export.json()["format"] == "pdf"

    exported = fitz.open(export.json()["export_file_uri"])
    try:
        text = exported[0].get_text()
    finally:
        exported.close()
    assert "2026-05-02" in text
    assert "शुल्क:" in text
    assert "RES-2026-004" in text


def test_detect_source_language_suggests_nepali_for_nepali_docx(client: TestClient) -> None:
    fixture = SAMPLES_DIR / "public-service-nepali-1.docx"
    with fixture.open("rb") as handle:
        response = client.post(
            "/api/documents/detect-source",
            files={
                "file": (
                    fixture.name,
                    handle,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )

    assert response.status_code == 200
    payload = response.json()
    assert payload["source_lang"] == "ne"
    assert payload["confidence"] in {"high", "medium"}


def _create_pdf(path: Path, lines: list[str]) -> None:
    document = fitz.open()
    page = document.new_page(width=595, height=842)
    y = 72
    for index, line in enumerate(lines):
        page.insert_textbox(
            fitz.Rect(56, y, 520, y + 28),
            line,
            fontname="helv",
            fontsize=16 if index == 0 else 12,
            color=(0, 0, 0),
        )
        y += 34 if index == 0 else 26
    document.save(path)
    document.close()


def test_process_and_get_document_include_trust_summary(client: TestClient) -> None:
    fixture = SAMPLES_DIR / "public-service-1.docx"
    with fixture.open("rb") as handle:
        upload = client.post(
            "/api/documents",
            data={"source_lang": "en", "target_lang": "ne", "domain": "public_service", "subdomain": "residence"},
            files={
                "file": (
                    fixture.name,
                    handle,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )

    document_id = upload.json()["id"]
    process = client.post(f"/api/documents/{document_id}/process")
    assert process.status_code == 200
    process_payload = process.json()
    trust_summary = process_payload["trust_summary"]

    assert trust_summary["total_segments"] == process_payload["counts"]["segments"] == 14
    assert trust_summary["approved_segments"] == process_payload["counts"]["approved"] == 0
    assert trust_summary["memory_reused_segments"] == process_payload["counts"]["memory_applied"] == 0
    assert trust_summary["protected_values_total"] > 0
    assert trust_summary["protected_values_preserved"] == trust_summary["protected_values_total"]
    assert trust_summary["unresolved_review_flags"] == process_payload["counts"]["needs_review"]

    document = client.get(f"/api/documents/{document_id}")
    assert document.status_code == 200
    assert document.json()["trust_summary"] == trust_summary


def test_feedback_pack_requires_all_segments_to_be_approved(client: TestClient) -> None:
    fixture = SAMPLES_DIR / "public-service-1.docx"
    with fixture.open("rb") as handle:
        upload = client.post(
            "/api/documents",
            data={"source_lang": "en", "target_lang": "ne", "domain": "public_service", "subdomain": "residence"},
            files={
                "file": (
                    fixture.name,
                    handle,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )

    document_id = upload.json()["id"]
    assert client.post(f"/api/documents/{document_id}/process").status_code == 200

    response = client.get(f"/api/documents/{document_id}/feedback-pack")

    assert response.status_code == 409
    assert "approved" in response.json()["detail"].lower()


def test_feedback_pack_contains_redacted_rows_and_manifest(client: TestClient) -> None:
    fixture = SAMPLES_DIR / "public-service-1.docx"
    with fixture.open("rb") as handle:
        upload = client.post(
            "/api/documents",
            data={"source_lang": "en", "target_lang": "ne", "domain": "public_service", "subdomain": "residence"},
            files={
                "file": (
                    fixture.name,
                    handle,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )

    document_id = upload.json()["id"]
    assert client.post(f"/api/documents/{document_id}/process").status_code == 200
    for segment in client.get(f"/api/documents/{document_id}/segments").json():
        response = client.post(
            f"/api/segments/{segment['id']}/approve",
            json={"text": segment["translation"]["candidate_text"], "actor": "test-reviewer"},
        )
        assert response.status_code == 200

    response = client.get(f"/api/documents/{document_id}/feedback-pack")
    assert response.status_code == 200

    pack = _read_feedback_pack(response.content)
    manifest = pack["manifest.json"]
    approved_rows = pack["approved_segments.tsv"]
    correction_rows = pack["review_corrections.tsv"]

    assert manifest["approved_segments"] == 14
    assert manifest["included_rows"] < manifest["approved_segments"]
    assert manifest["redaction_mode"] == "protected_entity_placeholders_v1"
    assert correction_rows == []
    assert any(row["source_text_redacted"] == "Fee: <MONEY>" for row in approved_rows)
    assert any("<URL>" in row["source_text_redacted"] for row in approved_rows)
    assert "+977" not in json.dumps(approved_rows, ensure_ascii=False)
    assert "RES-2026-005" not in json.dumps(approved_rows, ensure_ascii=False)
    assert all("Maya Lama" not in row["source_text_redacted"] for row in approved_rows)


def test_feedback_pack_exports_review_corrections(client: TestClient, tmp_path: Path) -> None:
    source = tmp_path / "fee-only.txt"
    source.write_text("Fee: NPR 500\n", encoding="utf-8")

    with source.open("rb") as handle:
        upload = client.post(
            "/api/documents",
            data={"source_lang": "en", "target_lang": "ne", "domain": "public_service", "subdomain": "residence"},
            files={"file": ("fee-only.txt", handle, "text/plain")},
        )

    document_id = upload.json()["id"]
    assert client.post(f"/api/documents/{document_id}/process").status_code == 200
    segment = client.get(f"/api/documents/{document_id}/segments").json()[0]
    assert client.patch(
        f"/api/segments/{segment['id']}/translation",
        json={"candidate_text": "शुल्क"},
    ).status_code == 200
    assert client.post(
        f"/api/segments/{segment['id']}/approve",
        json={"text": "शुल्क", "actor": "test-reviewer"},
    ).status_code == 200

    response = client.get(f"/api/documents/{document_id}/feedback-pack")
    assert response.status_code == 200

    pack = _read_feedback_pack(response.content)
    manifest = pack["manifest.json"]
    correction_rows = pack["review_corrections.tsv"]

    assert manifest["corrected_segments"] == 1
    assert len(correction_rows) == 1
    assert correction_rows[0]["source_text_redacted"] == "Fee: <MONEY>"
    assert correction_rows[0]["candidate_text_redacted"] == "शुल्क: <MONEY>"
    assert correction_rows[0]["approved_text_redacted"] == "शुल्क"
    assert correction_rows[0]["review_action"] == "revised_then_approve"


def test_feedback_pack_redacts_sensitive_values_introduced_during_review(client: TestClient, tmp_path: Path) -> None:
    source = tmp_path / "fee-only.txt"
    source.write_text("Fee: NPR 500\n", encoding="utf-8")

    with source.open("rb") as handle:
        upload = client.post(
            "/api/documents",
            data={"source_lang": "en", "target_lang": "ne", "domain": "public_service", "subdomain": "residence"},
            files={"file": ("fee-only.txt", handle, "text/plain")},
        )

    document_id = upload.json()["id"]
    assert client.post(f"/api/documents/{document_id}/process").status_code == 200
    segment = client.get(f"/api/documents/{document_id}/segments").json()[0]
    assert client.patch(
        f"/api/segments/{segment['id']}/translation",
        json={"candidate_text": "फोन: +977-9841234567"},
    ).status_code == 200
    assert client.post(
        f"/api/segments/{segment['id']}/approve",
        json={"text": "फोन: +977-9841234567", "actor": "test-reviewer"},
    ).status_code == 200

    response = client.get(f"/api/documents/{document_id}/feedback-pack")
    assert response.status_code == 200

    pack = _read_feedback_pack(response.content)
    correction_rows = pack["review_corrections.tsv"]

    assert len(correction_rows) == 1
    assert correction_rows[0]["approved_text_redacted"] == "फोन: <PHONE>"
    assert "+977-9841234567" not in json.dumps(pack, ensure_ascii=False)


def test_feedback_pack_manifest_tracks_memory_reuse(client: TestClient) -> None:
    fixture = SAMPLES_DIR / "public-service-1.docx"
    with fixture.open("rb") as handle:
        upload_one = client.post(
            "/api/documents",
            data={"source_lang": "en", "target_lang": "ne", "domain": "public_service", "subdomain": "residence"},
            files={
                "file": (
                    fixture.name,
                    handle,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )

    first_document_id = upload_one.json()["id"]
    assert client.post(f"/api/documents/{first_document_id}/process").status_code == 200
    for segment in client.get(f"/api/documents/{first_document_id}/segments").json():
        response = client.post(
            f"/api/segments/{segment['id']}/approve",
            json={"text": segment["translation"]["candidate_text"], "actor": "test-reviewer"},
        )
        assert response.status_code == 200

    fixture_two = SAMPLES_DIR / "public-service-2.docx"
    with fixture_two.open("rb") as handle:
        upload_two = client.post(
            "/api/documents",
            data={"source_lang": "en", "target_lang": "ne", "domain": "public_service", "subdomain": "residence"},
            files={
                "file": (
                    fixture_two.name,
                    handle,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )

    second_document_id = upload_two.json()["id"]
    assert client.post(f"/api/documents/{second_document_id}/process").status_code == 200
    client.post(f"/api/documents/{second_document_id}/approve-unflagged").raise_for_status()
    for segment in client.get(f"/api/documents/{second_document_id}/segments").json():
        if segment["status"] != "approved":
            response = client.post(
                f"/api/segments/{segment['id']}/approve",
                json={"text": segment["translation"]["candidate_text"], "actor": "test-reviewer"},
            )
            assert response.status_code == 200

    response = client.get(f"/api/documents/{second_document_id}/feedback-pack")
    assert response.status_code == 200

    manifest = _read_feedback_pack(response.content)["manifest.json"]
    assert manifest["memory_reused_segments"] > 0


def test_feedback_pack_supports_tamang_document_flow(client: TestClient) -> None:
    fixture = SAMPLES_DIR / "public-service-1.docx"
    with fixture.open("rb") as handle:
        upload = client.post(
            "/api/documents",
            data={"source_lang": "en", "target_lang": "tmg", "domain": "public_service", "subdomain": "residence"},
            files={
                "file": (
                    fixture.name,
                    handle,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )

    document_id = upload.json()["id"]
    assert client.post(f"/api/documents/{document_id}/process").status_code == 200
    for segment in client.get(f"/api/documents/{document_id}/segments").json():
        response = client.post(
            f"/api/segments/{segment['id']}/approve",
            json={"text": segment["translation"]["candidate_text"], "actor": "test-reviewer"},
        )
        assert response.status_code == 200

    response = client.get(f"/api/documents/{document_id}/feedback-pack")
    assert response.status_code == 200

    pack = _read_feedback_pack(response.content)
    assert pack["manifest.json"]["target_lang"] == "tmg"
    assert pack["approved_segments.tsv"]
    assert all("Maya Lama" not in row["source_text_redacted"] for row in pack["approved_segments.tsv"])


def test_upload_rejects_blank_domain(client: TestClient) -> None:
    response = client.post(
        "/api/documents",
        data={"source_lang": "en", "target_lang": "ne", "domain": "   ", "subdomain": "residence"},
        files={
            "file": (
                "public-service-1.docx",
                b"placeholder",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )

    assert response.status_code == 422
    assert response.json()["detail"] == "domain must not be empty."


def test_upload_rejects_unsupported_language_codes(client: TestClient) -> None:
    response = client.post(
        "/api/documents",
        data={"source_lang": "fr", "target_lang": "ne", "domain": "public_service", "subdomain": "residence"},
        files={
            "file": (
                "public-service-1.docx",
                b"placeholder",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )

    assert response.status_code == 422
    assert "Unsupported source_lang" in response.json()["detail"]


def test_upload_rejects_same_language_pair(client: TestClient) -> None:
    response = client.post(
        "/api/documents",
        data={"source_lang": "tmg", "target_lang": "tmg", "domain": "public_service", "subdomain": "residence"},
        files={
            "file": (
                "public-service-1.docx",
                b"placeholder",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )

    assert response.status_code == 422
    assert response.json()["detail"] == "source_lang and target_lang must be different."


def test_duplicate_detection_respects_language_pair(client: TestClient) -> None:
    fixture = SAMPLES_DIR / "public-service-1.docx"

    with fixture.open("rb") as handle:
        first = client.post(
            "/api/documents",
            data={"source_lang": "en", "target_lang": "ne", "domain": "public_service", "subdomain": "residence"},
            files={
                "file": (
                    fixture.name,
                    handle,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )

    assert first.status_code == 200
    first_payload = first.json()
    assert first_payload["is_duplicate"] is False

    with fixture.open("rb") as handle:
        second = client.post(
            "/api/documents",
            data={"source_lang": "en", "target_lang": "tmg", "domain": "public_service", "subdomain": "residence"},
            files={
                "file": (
                    fixture.name,
                    handle,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )

    assert second.status_code == 200
    second_payload = second.json()
    assert second_payload["is_duplicate"] is False
    assert second_payload["id"] != first_payload["id"]


def test_env_example_contains_bootable_settings(monkeypatch: pytest.MonkeyPatch) -> None:
    for key in [
        "SANAD_DATABASE_URL",
        "SANAD_STORAGE_ROOT",
        "SANAD_ACTIVE_PROVIDER",
        "SANAD_TMT_API_ENDPOINT",
        "SANAD_TMT_API_KEY",
        "SANAD_TMT_AUTH_METHOD",
        "SANAD_TMT_TIMEOUT_SECONDS",
        "SANAD_TMT_PROVIDER_BATCH_SIZE",
    ]:
        monkeypatch.delenv(key, raising=False)

    settings = Settings(_env_file=ROOT_DIR / ".env.example")

    assert settings.active_provider == "tmt_api"
    assert settings.tmt_api_endpoint == "https://tmt.ilprl.ku.edu.np"
    assert settings.tmt_timeout_seconds == 20
    assert settings.tmt_provider_batch_size == 25


def test_empty_docx_fails_with_parse_quality_message(client: TestClient, tmp_path: Path) -> None:
    empty_docx = tmp_path / "empty.docx"
    Document().save(empty_docx)

    with empty_docx.open("rb") as handle:
        upload = client.post(
            "/api/documents",
            data={"source_lang": "en", "target_lang": "ne", "domain": "public_service", "subdomain": "residence"},
            files={
                "file": (
                    "empty.docx",
                    handle,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )

    assert upload.status_code == 200
    document_id = upload.json()["id"]
    process = client.post(f"/api/documents/{document_id}/process")

    assert process.status_code == 400
    assert "No translatable text" in process.json()["detail"]


def test_debug_provider_handles_invalid_provider_name(client: TestClient, monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setenv("SANAD_ACTIVE_PROVIDER", "broken-provider")
    get_settings.cache_clear()

    app = create_app()
    app.dependency_overrides[get_db] = client.app.dependency_overrides[get_db]
    with TestClient(app) as invalid_client:
        response = invalid_client.get("/api/debug/provider")

    assert response.status_code == 200
    payload = response.json()
    assert payload["active_provider"] == "broken-provider"
    assert payload["implemented"] is False
    assert "Unsupported SANAD_ACTIVE_PROVIDER" in payload["notes"]
    get_settings.cache_clear()


def test_editing_approved_segment_clears_memory_link(client: TestClient) -> None:
    fixture = SAMPLES_DIR / "public-service-1.docx"
    with fixture.open("rb") as handle:
        upload = client.post(
            "/api/documents",
            data={"source_lang": "en", "target_lang": "ne", "domain": "public_service", "subdomain": "residence"},
            files={
                "file": (
                    fixture.name,
                    handle,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )

    document_id = upload.json()["id"]
    process = client.post(f"/api/documents/{document_id}/process")
    assert process.status_code == 200

    segments = client.get(f"/api/documents/{document_id}/segments").json()
    first_segment = segments[0]
    approve = client.post(
        f"/api/segments/{first_segment['id']}/approve",
        json={"text": first_segment["translation"]["candidate_text"], "actor": "test-reviewer"},
    )
    assert approve.status_code == 200
    assert approve.json()["translation"]["memory_entry_id"] is not None

    edited = client.patch(
        f"/api/segments/{first_segment['id']}/translation",
        json={"candidate_text": "बसोबास प्रमाणपत्र अनुरोध (edited)"},
    )
    assert edited.status_code == 200
    payload = edited.json()
    assert payload["status"] == "needs_review"
    assert payload["translation"]["approved_text"] is None
    assert payload["translation"]["memory_entry_id"] is None


def test_blank_translation_edit_is_rejected(client: TestClient) -> None:
    fixture = SAMPLES_DIR / "public-service-1.docx"
    with fixture.open("rb") as handle:
        upload = client.post(
            "/api/documents",
            data={"source_lang": "en", "target_lang": "ne", "domain": "public_service", "subdomain": "residence"},
            files={
                "file": (
                    fixture.name,
                    handle,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )

    document_id = upload.json()["id"]
    assert client.post(f"/api/documents/{document_id}/process").status_code == 200
    segment = client.get(f"/api/documents/{document_id}/segments").json()[0]

    response = client.patch(
        f"/api/segments/{segment['id']}/translation",
        json={"candidate_text": "   "},
    )

    assert response.status_code == 422
    assert "candidate_text must not be empty" in response.text


def test_blank_approval_text_is_rejected(client: TestClient) -> None:
    fixture = SAMPLES_DIR / "public-service-1.docx"
    with fixture.open("rb") as handle:
        upload = client.post(
            "/api/documents",
            data={"source_lang": "en", "target_lang": "ne", "domain": "public_service", "subdomain": "residence"},
            files={
                "file": (
                    fixture.name,
                    handle,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )

    document_id = upload.json()["id"]
    assert client.post(f"/api/documents/{document_id}/process").status_code == 200
    segment = client.get(f"/api/documents/{document_id}/segments").json()[0]

    response = client.post(
        f"/api/segments/{segment['id']}/approve",
        json={"text": "   ", "actor": "test-reviewer"},
    )

    assert response.status_code == 422
    assert "text must not be empty" in response.text


def test_translation_edit_recomputes_risk_flags(client: TestClient) -> None:
    fixture = SAMPLES_DIR / "public-service-1.docx"
    with fixture.open("rb") as handle:
        upload = client.post(
            "/api/documents",
            data={"source_lang": "en", "target_lang": "ne", "domain": "public_service", "subdomain": "residence"},
            files={
                "file": (
                    fixture.name,
                    handle,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )

    document_id = upload.json()["id"]
    assert client.post(f"/api/documents/{document_id}/process").status_code == 200
    segments = client.get(f"/api/documents/{document_id}/segments").json()
    fee_segment = next(segment for segment in segments if segment["source_text"] == "Fee: NPR 500")

    response = client.patch(
        f"/api/segments/{fee_segment['id']}/translation",
        json={"candidate_text": "शुल्क"},
    )

    assert response.status_code == 200
    payload = response.json()
    reason_codes = {reason["code"] for reason in payload["translation"]["risk_reasons"]}
    assert payload["status"] == "needs_review"
    assert payload["translation"]["status"] == "needs_review"
    assert "changed_number" in reason_codes
    assert "changed_protected_entity" in reason_codes


def test_memory_hit_includes_provenance_fields(client: TestClient) -> None:
    fixture = SAMPLES_DIR / "public-service-1.docx"
    with fixture.open("rb") as handle:
        upload_one = client.post(
            "/api/documents",
            data={"source_lang": "en", "target_lang": "ne", "domain": "public_service", "subdomain": "residence"},
            files={
                "file": (
                    fixture.name,
                    handle,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )

    first_document_id = upload_one.json()["id"]
    assert client.post(f"/api/documents/{first_document_id}/process").status_code == 200
    for segment in client.get(f"/api/documents/{first_document_id}/segments").json():
        response = client.post(
            f"/api/segments/{segment['id']}/approve",
            json={"text": segment["translation"]["candidate_text"], "actor": "test-reviewer"},
        )
        assert response.status_code == 200

    fixture_two = SAMPLES_DIR / "public-service-2.docx"
    with fixture_two.open("rb") as handle:
        upload_two = client.post(
            "/api/documents",
            data={"source_lang": "en", "target_lang": "ne", "domain": "public_service", "subdomain": "residence"},
            files={
                "file": (
                    fixture_two.name,
                    handle,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )

    second_document_id = upload_two.json()["id"]
    process_two = client.post(f"/api/documents/{second_document_id}/process")
    assert process_two.status_code == 200
    assert process_two.json()["counts"]["memory_applied"] > 0
    assert process_two.json()["trust_summary"]["memory_reused_segments"] == process_two.json()["counts"]["memory_applied"]

    segments = client.get(f"/api/documents/{second_document_id}/segments").json()
    memory_hit = next(segment for segment in segments if segment["translation"]["source_type"] == "memory")
    provenance = memory_hit["translation"]["memory_provenance"]

    assert provenance["source_document_filename"] == "public-service-1.docx"
    assert provenance["scope_label"] == "Public Service / Residence"
    assert provenance["approved_at"]
    assert provenance["times_used"] >= 1


def test_same_text_in_different_subdomain_does_not_reuse_memory(client: TestClient) -> None:
    fixture = SAMPLES_DIR / "public-service-1.docx"
    with fixture.open("rb") as handle:
        upload_one = client.post(
            "/api/documents",
            data={"source_lang": "en", "target_lang": "ne", "domain": "public_service", "subdomain": "residence"},
            files={
                "file": (
                    fixture.name,
                    handle,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )

    first_document_id = upload_one.json()["id"]
    assert client.post(f"/api/documents/{first_document_id}/process").status_code == 200
    for segment in client.get(f"/api/documents/{first_document_id}/segments").json():
        response = client.post(
            f"/api/segments/{segment['id']}/approve",
            json={"text": segment["translation"]["candidate_text"], "actor": "test-reviewer"},
        )
        assert response.status_code == 200

    proof_fixture = SAMPLES_DIR / "public-service-cross-scope-proof.docx"
    with proof_fixture.open("rb") as handle:
        upload_two = client.post(
            "/api/documents",
            data={"source_lang": "en", "target_lang": "ne", "domain": "public_service", "subdomain": "citizenship"},
            files={
                "file": (
                    proof_fixture.name,
                    handle,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )

    second_document_id = upload_two.json()["id"]
    process_two = client.post(f"/api/documents/{second_document_id}/process")
    assert process_two.status_code == 200
    assert process_two.json()["counts"]["memory_applied"] == 0


def test_edit_and_approve_recomputes_risk_before_memory_write(client: TestClient) -> None:
    fixture = SAMPLES_DIR / "public-service-1.docx"
    with fixture.open("rb") as handle:
        upload = client.post(
            "/api/documents",
            data={"source_lang": "en", "target_lang": "ne", "domain": "public_service", "subdomain": "residence"},
            files={
                "file": (
                    fixture.name,
                    handle,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )

    document_id = upload.json()["id"]
    assert client.post(f"/api/documents/{document_id}/process").status_code == 200
    segments = client.get(f"/api/documents/{document_id}/segments").json()
    fee_segment = next(segment for segment in segments if segment["source_text"] == "Fee: NPR 500")

    response = client.post(
        f"/api/segments/{fee_segment['id']}/approve",
        json={"text": "शुल्क", "actor": "test-reviewer"},
    )

    assert response.status_code == 200
    payload = response.json()
    reason_codes = {reason["code"] for reason in payload["translation"]["risk_reasons"]}
    assert payload["status"] == "approved"
    assert payload["translation"]["memory_entry_id"] is not None
    assert "changed_number" in reason_codes
    assert "changed_protected_entity" in reason_codes


def test_trust_summary_tracks_protected_value_regression_and_manual_resolution(client: TestClient, tmp_path: Path) -> None:
    source = tmp_path / "fee-only.txt"
    source.write_text("Fee: NPR 500\n", encoding="utf-8")

    with source.open("rb") as handle:
        upload = client.post(
            "/api/documents",
            data={"source_lang": "en", "target_lang": "ne", "domain": "public_service", "subdomain": "residence"},
            files={"file": ("fee-only.txt", handle, "text/plain")},
        )

    document_id = upload.json()["id"]
    process = client.post(f"/api/documents/{document_id}/process")
    assert process.status_code == 200
    initial_summary = process.json()["trust_summary"]
    assert initial_summary == {
        "total_segments": 1,
        "approved_segments": 0,
        "memory_reused_segments": 0,
        "protected_values_total": 1,
        "protected_values_preserved": 1,
        "unresolved_review_flags": 0,
    }

    segment = client.get(f"/api/documents/{document_id}/segments").json()[0]
    edited = client.patch(
        f"/api/segments/{segment['id']}/translation",
        json={"candidate_text": "शुल्क"},
    )
    assert edited.status_code == 200

    after_edit = client.get(f"/api/documents/{document_id}").json()["trust_summary"]
    assert after_edit["protected_values_total"] == 1
    assert after_edit["protected_values_preserved"] == 0
    assert after_edit["unresolved_review_flags"] == 1
    assert after_edit["approved_segments"] == 0

    approved = client.post(
        f"/api/segments/{segment['id']}/approve",
        json={"text": "शुल्क", "actor": "test-reviewer"},
    )
    assert approved.status_code == 200

    after_approval = client.get(f"/api/documents/{document_id}").json()["trust_summary"]
    assert after_approval["protected_values_total"] == 1
    assert after_approval["protected_values_preserved"] == 0
    assert after_approval["unresolved_review_flags"] == 0
    assert after_approval["approved_segments"] == 1


def test_reprocess_clears_old_review_rows_without_orphans(client: TestClient) -> None:
    fixture = SAMPLES_DIR / "public-service-1.docx"
    with fixture.open("rb") as handle:
        upload = client.post(
            "/api/documents",
            data={"source_lang": "en", "target_lang": "ne", "domain": "public_service", "subdomain": "residence"},
            files={
                "file": (
                    fixture.name,
                    handle,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )

    document_id = upload.json()["id"]
    assert client.post(f"/api/documents/{document_id}/process").status_code == 200
    segments = client.get(f"/api/documents/{document_id}/segments").json()
    first_segment = segments[0]
    assert client.post(
        f"/api/segments/{first_segment['id']}/approve",
        json={"text": first_segment["translation"]["candidate_text"], "actor": "test-reviewer"},
    ).status_code == 200

    assert client.post(f"/api/documents/{document_id}/process").status_code == 200
    refreshed = client.get(f"/api/documents/{document_id}").json()
    assert refreshed["status"] == "processed"
    assert refreshed["export_file_uri"] is None

    with client.app.state.session_local() as db:
        from sanad_api.models import ReviewEvent, Segment, Translation

        assert db.query(Segment).count() == 14
        assert db.query(Translation).count() == 14
        assert db.query(ReviewEvent).count() == 0


def test_failed_reprocess_keeps_previous_review_state(client: TestClient, monkeypatch: pytest.MonkeyPatch) -> None:
    fixture = SAMPLES_DIR / "public-service-1.docx"
    with fixture.open("rb") as handle:
        upload = client.post(
            "/api/documents",
            data={"source_lang": "en", "target_lang": "ne", "domain": "public_service", "subdomain": "residence"},
            files={
                "file": (
                    fixture.name,
                    handle,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )

    document_id = upload.json()["id"]
    assert client.post(f"/api/documents/{document_id}/process").status_code == 200
    original_segments = client.get(f"/api/documents/{document_id}/segments").json()
    assert original_segments

    monkeypatch.setenv("SANAD_ACTIVE_PROVIDER", "broken-provider")
    get_settings.cache_clear()
    failed = client.post(f"/api/documents/{document_id}/process")
    assert failed.status_code == 400
    assert "Translation provider failed before review" in failed.json()["detail"]

    after_failure = client.get(f"/api/documents/{document_id}/segments").json()
    assert len(after_failure) == len(original_segments)
    assert [segment["source_text"] for segment in after_failure] == [segment["source_text"] for segment in original_segments]
    get_settings.cache_clear()


@pytest.mark.parametrize(
    ("source_lang", "target_lang", "doc1_name", "doc2_name"),
    [
        ("en", "ne", "public-service-1.docx", "public-service-2.docx"),
        ("en", "tmg", "public-service-1.docx", "public-service-2.docx"),
        ("ne", "en", "public-service-nepali-1.docx", "public-service-nepali-2.docx"),
        ("ne", "tmg", "public-service-nepali-1.docx", "public-service-nepali-2.docx"),
        ("tmg", "en", "public-service-tamang-1.docx", "public-service-tamang-2.docx"),
        ("tmg", "ne", "public-service-tamang-1.docx", "public-service-tamang-2.docx"),
    ],
)
def test_multilingual_fixture_flow_supports_memory_reuse(
    client: TestClient,
    source_lang: str,
    target_lang: str,
    doc1_name: str,
    doc2_name: str,
) -> None:
    fixture_one = SAMPLES_DIR / doc1_name
    with fixture_one.open("rb") as handle:
        upload_one = client.post(
            "/api/documents",
            data={"source_lang": source_lang, "target_lang": target_lang, "domain": "public_service", "subdomain": "residence"},
            files={
                "file": (
                    fixture_one.name,
                    handle,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )

    first_document_id = upload_one.json()["id"]
    assert client.post(f"/api/documents/{first_document_id}/process").status_code == 200
    for segment in client.get(f"/api/documents/{first_document_id}/segments").json():
        response = client.post(
            f"/api/segments/{segment['id']}/approve",
            json={"text": segment["translation"]["candidate_text"], "actor": "test-reviewer"},
        )
        assert response.status_code == 200

    fixture_two = SAMPLES_DIR / doc2_name
    with fixture_two.open("rb") as handle:
        upload_two = client.post(
            "/api/documents",
            data={"source_lang": source_lang, "target_lang": target_lang, "domain": "public_service", "subdomain": "residence"},
            files={
                "file": (
                    fixture_two.name,
                    handle,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )

    second_document_id = upload_two.json()["id"]
    process_two = client.post(f"/api/documents/{second_document_id}/process")
    assert process_two.status_code == 200
    assert process_two.json()["counts"]["memory_applied"] > 0


def _read_feedback_pack(content: bytes) -> dict[str, object]:
    with zipfile.ZipFile(io.BytesIO(content)) as archive:
        manifest = json.loads(archive.read("manifest.json").decode("utf-8"))
        approved_rows = list(csv.DictReader(io.StringIO(archive.read("approved_segments.tsv").decode("utf-8")), delimiter="\t"))
        correction_rows = list(
            csv.DictReader(io.StringIO(archive.read("review_corrections.tsv").decode("utf-8")), delimiter="\t")
        )
    return {
        "manifest.json": manifest,
        "approved_segments.tsv": approved_rows,
        "review_corrections.tsv": correction_rows,
    }
