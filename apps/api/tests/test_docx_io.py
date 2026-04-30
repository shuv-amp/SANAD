from pathlib import Path

from docx import Document

from sanad_api.services.docx_io import export_docx, parse_docx


def test_parse_and_export_docx_round_trip(tmp_path: Path) -> None:
    source = tmp_path / "source.docx"
    output = tmp_path / "translated.docx"

    document = Document()
    document.add_paragraph("Certificate of Residence Request")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Phone"
    table.rows[0].cells[1].text = "+977-9841234567"
    document.save(source)

    segments = parse_docx(source)

    assert [segment.source_text for segment in segments] == [
        "Certificate of Residence Request",
        "Phone",
        "+977-9841234567",
    ]

    export_docx(
        source,
        [
            (segments[0].location_json, "बसोबास प्रमाणपत्र अनुरोध"),
            (segments[1].location_json, "फोन"),
            (segments[2].location_json, "+977-9841234567"),
        ],
        output,
    )

    exported = Document(output)
    assert exported.paragraphs[0].text == "बसोबास प्रमाणपत्र अनुरोध"
    assert exported.tables[0].rows[0].cells[0].text == "फोन"
    assert exported.tables[0].rows[0].cells[1].text == "+977-9841234567"

