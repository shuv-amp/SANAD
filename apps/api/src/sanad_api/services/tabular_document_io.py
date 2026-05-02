import csv
from pathlib import Path

from sanad_api.services.docx_io import ParsedSegment
from sanad_api.services.normalization import contains_devanagari, display_normalize


TABULAR_DOCUMENT_TYPES = {"csv": ",", "tsv": "\t"}


def parse_tabular_document(path: Path, file_type: str) -> list[ParsedSegment]:
    delimiter = TABULAR_DOCUMENT_TYPES[file_type]
    rows = _read_rows(path, delimiter)
    segments: list[ParsedSegment] = []
    sequence = 0

    for row_index, row in enumerate(rows):
        for cell_index, value in enumerate(row):
            text = display_normalize(value)
            if not text:
                continue
            sequence += 1
            segments.append(
                ParsedSegment(
                    sequence=sequence,
                    segment_type="table_cell",
                    source_text=text,
                    location_json={"kind": "table_cell", "row_index": row_index, "cell_index": cell_index},
                )
            )

    return segments


def export_tabular_document(
    original_path: Path,
    translations_by_location: list[tuple[dict[str, int], str]],
    output_path: Path,
    file_type: str,
) -> Path:
    delimiter = TABULAR_DOCUMENT_TYPES[file_type]
    rows = _read_rows(original_path, delimiter)
    replacements = {(location["row_index"], location["cell_index"]): text for location, text in translations_by_location}

    for row_index, row in enumerate(rows):
        for cell_index, _ in enumerate(row):
            replacement = replacements.get((row_index, cell_index))
            if replacement is not None:
                row[cell_index] = replacement

    output_path.parent.mkdir(parents=True, exist_ok=True)
    with output_path.open("w", encoding="utf-8", newline="") as handle:
        writer = csv.writer(handle, delimiter=delimiter)
        writer.writerows(rows)
    return output_path


def export_tabular_docx(
    original_path: Path,
    translations_by_location: list[tuple[dict[str, int], str]],
    output_path: Path,
    file_type: str,
) -> Path:
    """Generate a Word document containing a table that matches the structure of the source tabular document."""
    from docx import Document as DocxDocument
    from sanad_api.services.text_document_io import _apply_script_font

    delimiter = TABULAR_DOCUMENT_TYPES[file_type]
    rows = _read_rows(original_path, delimiter)
    replacements = {(location["row_index"], location["cell_index"]): text for location, text in translations_by_location}

    doc = DocxDocument()
    if not rows:
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))
        return output_path

    # Determine table dimensions
    max_cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.style = "Table Grid"

    for r_idx, row in enumerate(rows):
        for c_idx, _ in enumerate(row):
            text = replacements.get((r_idx, c_idx), row[c_idx])
            cell = table.cell(r_idx, c_idx)
            cell.text = text
            # Ensure the correct font is applied to all text in the cell
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    _apply_script_font(run, text)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


def _read_rows(path: Path, delimiter: str) -> list[list[str]]:
    with path.open("r", encoding="utf-8", errors="replace", newline="") as handle:
        reader = csv.reader(handle, delimiter=delimiter)
        return [list(row) for row in reader]

