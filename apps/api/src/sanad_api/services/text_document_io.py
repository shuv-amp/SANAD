import re
import subprocess
from pathlib import Path

from docx import Document as DocxDocument
from docx.oxml.ns import qn
from docx.text.run import Run

from sanad_api.services.docx_io import ParsedSegment
from sanad_api.services.normalization import contains_devanagari, display_normalize


DIRECT_TEXT_TYPES = {"txt", "md"}
TEXTUTIL_TEXT_TYPES = {"doc", "odt", "rtf", "html"}
TEXT_DOCUMENT_TYPES = DIRECT_TEXT_TYPES | TEXTUTIL_TEXT_TYPES


def parse_text_document(path: Path, file_type: str) -> list[ParsedSegment]:
    plain_text = _extract_plain_text(path, file_type)
    blocks = _split_text_blocks(plain_text)
    return [
        ParsedSegment(
            sequence=index,
            segment_type="text_block",
            source_text=block,
            location_json={"kind": "text_block", "block_index": index - 1},
        )
        for index, block in enumerate(blocks, start=1)
    ]


def export_text_docx(translated_blocks: list[str], output_path: Path) -> Path:
    document = DocxDocument()
    for block in translated_blocks:
        paragraph = document.add_paragraph()
        run = paragraph.add_run(block)
        _apply_script_font(run, block)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))
    return output_path


def _extract_plain_text(path: Path, file_type: str) -> str:
    normalized_type = file_type.lower()
    if normalized_type in DIRECT_TEXT_TYPES:
        return path.read_text(encoding="utf-8", errors="replace")
    if normalized_type in TEXTUTIL_TEXT_TYPES:
        try:
            result = subprocess.run(
                ["/usr/bin/textutil", "-convert", "txt", "-stdout", str(path)],
                check=True,
                capture_output=True,
                text=True,
                encoding="utf-8",
                errors="replace",
            )
        except FileNotFoundError as exc:
            raise ValueError("This machine is missing textutil, so SANAD cannot read this text document format.") from exc
        except subprocess.CalledProcessError as exc:
            stderr = (exc.stderr or "").strip()
            if stderr:
                raise ValueError(f"SANAD could not read this {normalized_type.upper()} document: {stderr}") from exc
            raise ValueError(f"SANAD could not read this {normalized_type.upper()} document.") from exc
        return result.stdout
    raise ValueError(f"Unsupported text document type: {file_type}")


def _split_text_blocks(text: str) -> list[str]:
    normalized = text.replace("\r\n", "\n").replace("\r", "\n")
    raw_blocks = re.split(r"\n\s*\n+", normalized)
    blocks: list[str] = []
    for raw_block in raw_blocks:
        lines = [display_normalize(line) for line in raw_block.split("\n")]
        compact = " ".join(line.strip() for line in lines if line.strip())
        compact = display_normalize(compact)
        if compact:
            blocks.append(compact)
    return blocks


def _apply_script_font(run: Run, text: str) -> None:
    if not contains_devanagari(text):
        return
    font_name = "Noto Sans Devanagari"
    run.font.name = font_name
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = r_pr._add_rFonts()
    for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
        r_fonts.set(qn(f"w:{attr}"), font_name)
