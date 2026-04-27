from dataclasses import dataclass
from pathlib import Path
from typing import Any

from docx import Document as DocxDocument
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from sanad_api.services.normalization import contains_devanagari, display_normalize


@dataclass(frozen=True)
class ParsedSegment:
    sequence: int
    segment_type: str
    source_text: str
    location_json: dict[str, Any]


def parse_docx(path: Path) -> list[ParsedSegment]:
    document = DocxDocument(str(path))
    segments: list[ParsedSegment] = []
    sequence = 0

    for block_index, block in enumerate(document.iter_inner_content()):
        if isinstance(block, Paragraph):
            text = display_normalize(block.text)
            if text:
                sequence += 1
                segments.append(
                    ParsedSegment(
                        sequence=sequence,
                        segment_type="paragraph",
                        source_text=text,
                        location_json={"kind": "paragraph", "block_index": block_index},
                    )
                )
        elif isinstance(block, Table):
            for row_index, row in enumerate(block.rows):
                for cell_index, cell in enumerate(row.cells):
                    for paragraph_index, paragraph in enumerate(cell.paragraphs):
                        text = display_normalize(paragraph.text)
                        if text:
                            sequence += 1
                            segments.append(
                                ParsedSegment(
                                    sequence=sequence,
                                    segment_type="table_cell_paragraph",
                                    source_text=text,
                                    location_json={
                                        "kind": "table_cell_paragraph",
                                        "block_index": block_index,
                                        "row_index": row_index,
                                        "cell_index": cell_index,
                                        "paragraph_index": paragraph_index,
                                    },
                                )
                            )

    return segments


def export_docx(original_path: Path, translations_by_location: list[tuple[dict[str, Any], str]], output_path: Path) -> Path:
    document = DocxDocument(str(original_path))
    blocks = list(document.iter_inner_content())

    for location, translated_text in translations_by_location:
        paragraph = _paragraph_for_location(blocks, location)
        _replace_paragraph_text(paragraph, translated_text)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))
    return output_path


def _paragraph_for_location(blocks: list[Any], location: dict[str, Any]) -> Paragraph:
    block = blocks[location["block_index"]]
    if location["kind"] == "paragraph":
        if not isinstance(block, Paragraph):
            raise ValueError("DOCX location no longer points to a paragraph")
        return block

    if location["kind"] == "table_cell_paragraph":
        if not isinstance(block, Table):
            raise ValueError("DOCX location no longer points to a table")
        cell = block.rows[location["row_index"]].cells[location["cell_index"]]
        return cell.paragraphs[location["paragraph_index"]]

    raise ValueError(f"Unsupported DOCX location kind: {location['kind']}")


def _replace_paragraph_text(paragraph: Paragraph, text: str) -> None:
    if paragraph.runs:
        first_run = paragraph.runs[0]
        first_run.text = text
        _apply_script_font(first_run, text)
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        run = paragraph.add_run(text)
        _apply_script_font(run, text)


def _apply_script_font(run: Run, text: str) -> None:
    if not contains_devanagari(text):
        return
    font_name = "Noto Sans Devanagari"
    run.font.name = font_name
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = r_pr._add_rFonts()
    for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
        r_fonts.set(qn(f"w:{attr}"), font_name)

