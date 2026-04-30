import json
import os
import shutil
import zipfile
from pathlib import Path

from docx import Document

from sanad_api.services.docx_io import parse_docx
from sanad_api.services.demo_content import demo_text
from sanad_api.services.normalization import normalize_text


ROOT = Path(__file__).resolve().parents[3]
VALIDATION_DIR = ROOT / "apps" / "api" / "storage" / "language-coverage-validation"
DEVANAGARI_FONT = "Noto Sans Devanagari"
CASES = [
    {
        "name": "english_to_nepali",
        "source_lang": "en",
        "target_lang": "ne",
        "doc1": ROOT / "samples" / "demo" / "public-service-1.docx",
        "doc2": ROOT / "samples" / "demo" / "public-service-2.docx",
        "expected_title": demo_text("title", "ne"),
    },
    {
        "name": "english_to_tamang",
        "source_lang": "en",
        "target_lang": "tmg",
        "doc1": ROOT / "samples" / "demo" / "public-service-1.docx",
        "doc2": ROOT / "samples" / "demo" / "public-service-2.docx",
        "expected_title": demo_text("title", "tmg"),
    },
    {
        "name": "nepali_to_english",
        "source_lang": "ne",
        "target_lang": "en",
        "doc1": ROOT / "samples" / "demo" / "public-service-nepali-1.docx",
        "doc2": ROOT / "samples" / "demo" / "public-service-nepali-2.docx",
        "expected_title": demo_text("title", "en"),
    },
    {
        "name": "nepali_to_tamang",
        "source_lang": "ne",
        "target_lang": "tmg",
        "doc1": ROOT / "samples" / "demo" / "public-service-nepali-1.docx",
        "doc2": ROOT / "samples" / "demo" / "public-service-nepali-2.docx",
        "expected_title": demo_text("title", "tmg"),
    },
    {
        "name": "tamang_to_english",
        "source_lang": "tmg",
        "target_lang": "en",
        "doc1": ROOT / "samples" / "demo" / "public-service-tamang-1.docx",
        "doc2": ROOT / "samples" / "demo" / "public-service-tamang-2.docx",
        "expected_title": demo_text("title", "en"),
    },
    {
        "name": "tamang_to_nepali",
        "source_lang": "tmg",
        "target_lang": "ne",
        "doc1": ROOT / "samples" / "demo" / "public-service-tamang-1.docx",
        "doc2": ROOT / "samples" / "demo" / "public-service-tamang-2.docx",
        "expected_title": demo_text("title", "ne"),
    },
]


def main() -> None:
    if VALIDATION_DIR.exists():
        shutil.rmtree(VALIDATION_DIR)
    VALIDATION_DIR.mkdir(parents=True, exist_ok=True)

    os.environ["SANAD_DATABASE_URL"] = f"sqlite:///{VALIDATION_DIR / 'coverage.db'}"
    os.environ["SANAD_STORAGE_ROOT"] = str(VALIDATION_DIR / "files")
    os.environ["SANAD_ACTIVE_PROVIDER"] = "fixture"

    from fastapi.testclient import TestClient

    from sanad_api.config import get_settings
    from sanad_api.main import create_app

    get_settings.cache_clear()
    results: dict[str, dict] = {}
    with TestClient(create_app()) as client:
        for case in CASES:
            results[case["name"]] = _run_case(client, case)

    _assert_language_coverage(results)
    print(json.dumps(results, ensure_ascii=False, indent=2))
    print(f"Persisted multilingual validation outputs under: {VALIDATION_DIR}")


def _run_case(client, case: dict) -> dict:
    first = _process_document(client, case["doc1"], source_lang=case["source_lang"], target_lang=case["target_lang"])
    second = _process_document(client, case["doc2"], source_lang=case["source_lang"], target_lang=case["target_lang"])
    return {
        "source_lang": case["source_lang"],
        "target_lang": case["target_lang"],
        "doc1": first,
        "doc2": second,
        "expected_title": case["expected_title"],
    }


def _process_document(client, fixture: Path, *, source_lang: str, target_lang: str) -> dict:
    with fixture.open("rb") as handle:
        upload = client.post(
            "/api/documents",
            data={"source_lang": source_lang, "target_lang": target_lang, "domain": "public_service", "subdomain": "residence"},
            files={"file": (fixture.name, handle, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        )
    upload.raise_for_status()
    document_id = upload.json()["id"]

    process = client.post(f"/api/documents/{document_id}/process")
    process.raise_for_status()
    summary = process.json()
    segments = client.get(f"/api/documents/{document_id}/segments")
    segments.raise_for_status()
    segment_payload = segments.json()

    for segment in segment_payload:
        text = segment["translation"]["candidate_text"]
        response = client.post(f"/api/segments/{segment['id']}/approve", json={"text": text, "actor": "coverage-validator"})
        response.raise_for_status()

    export = client.post(f"/api/documents/{document_id}/export", json={"format": "docx"})
    export.raise_for_status()
    export_path = Path(export.json()["export_file_uri"])
    exported_text = _read_docx_text(export_path)
    ordered_text = [segment.source_text for segment in parse_docx(export_path)]
    return {
        "document_id": document_id,
        "segments": len(segment_payload),
        "memory_applied": summary["counts"]["memory_applied"],
        "export_exists": export_path.exists(),
        "export_path": str(export_path),
        "has_devanagari": _has_devanagari(exported_text),
        "declares_devanagari_font": _docx_declares_font(export_path, DEVANAGARI_FONT),
        "exported_text": exported_text,
        "ordered_text_preview": ordered_text[:6],
        "exported_text_preview": exported_text[:10],
    }


def _read_docx_text(path: Path) -> list[str]:
    document = Document(path)
    lines = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                lines.extend(paragraph.text for paragraph in cell.paragraphs if paragraph.text.strip())
    return lines


def _has_devanagari(lines: list[str]) -> bool:
    return any("\u0900" <= char <= "\u097F" for line in lines for char in line)


def _docx_declares_font(path: Path, font_name: str) -> bool:
    with zipfile.ZipFile(path) as archive:
        xml = archive.read("word/document.xml").decode("utf-8")
    return font_name in xml


def _assert_language_coverage(results: dict[str, dict]) -> None:
    for name, result in results.items():
        assert result["doc1"]["segments"] > 0, name
        assert result["doc2"]["segments"] > 0, name
        assert result["doc1"]["export_exists"], name
        assert result["doc2"]["export_exists"], name
        assert result["doc2"]["memory_applied"] > 0, name
        expected_title = normalize_text(result["expected_title"])
        doc1_lines = {normalize_text(line) for line in result["doc1"]["exported_text"]}
        doc2_lines = {normalize_text(line) for line in result["doc2"]["exported_text"]}
        assert expected_title in doc1_lines, name
        assert expected_title in doc2_lines, name
        if result["target_lang"] in {"ne", "tmg"}:
            assert result["doc1"]["has_devanagari"], name
            assert result["doc2"]["has_devanagari"], name


if __name__ == "__main__":
    main()
