import json
import io
import os
import shutil
import zipfile
from pathlib import Path

from docx import Document

from sanad_api.services.docx_io import parse_docx


ROOT = Path(__file__).resolve().parents[3]
FIXTURES = {
    "public-service-1": ROOT / "samples" / "demo" / "public-service-1.docx",
    "public-service-2": ROOT / "samples" / "demo" / "public-service-2.docx",
    "public-service-cross-scope-proof": ROOT / "samples" / "demo" / "public-service-cross-scope-proof.docx",
    "font-validation": ROOT / "samples" / "demo" / "font-validation.docx",
}
DEVANAGARI_FONT = "Noto Sans Devanagari"
VALIDATION_DIR = ROOT / "apps" / "api" / "storage" / "demo-validation"


def main() -> None:
    if VALIDATION_DIR.exists():
        shutil.rmtree(VALIDATION_DIR)
    VALIDATION_DIR.mkdir(parents=True, exist_ok=True)

    os.environ["SANAD_DATABASE_URL"] = f"sqlite:///{VALIDATION_DIR / 'demo.db'}"
    os.environ["SANAD_STORAGE_ROOT"] = str(VALIDATION_DIR / "files")
    os.environ["SANAD_ACTIVE_PROVIDER"] = "fixture"

    from fastapi.testclient import TestClient

    from sanad_api.config import get_settings
    from sanad_api.main import create_app

    get_settings.cache_clear()
    results: dict[str, dict] = {}
    with TestClient(create_app()) as client:
        first = _run_docx(client, "public-service-1", approve_all=True)
        second = _run_docx(client, "public-service-2", approve_all=False)
        font = _run_docx(client, "font-validation", approve_all=True)
        results["public-service-1"] = first
        results["public-service-2"] = second
        results["font-validation"] = font

    _assert_demo_expectations(results)
    print(json.dumps(results, ensure_ascii=False, indent=2))
    print(f"Persisted validation outputs under: {VALIDATION_DIR}")


def _run_docx(client, fixture_name: str, *, approve_all: bool) -> dict:
    fixture = FIXTURES[fixture_name]
    with fixture.open("rb") as handle:
        upload = client.post(
            "/api/documents",
            data={"source_lang": "en", "target_lang": "ne", "domain": "public_service", "subdomain": "residence"},
            files={
                "file": (
                    fixture.name,
                    handle,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )
    upload.raise_for_status()
    document_id = upload.json()["id"]

    process = client.post(f"/api/documents/{document_id}/process")
    process.raise_for_status()
    process_summary = process.json()
    segments = client.get(f"/api/documents/{document_id}/segments")
    segments.raise_for_status()
    segment_payload = segments.json()

    if approve_all:
        for segment in segment_payload:
            text = segment["translation"]["candidate_text"]
            response = client.post(f"/api/segments/{segment['id']}/approve", json={"text": text, "actor": "demo-validator"})
            response.raise_for_status()
    else:
        client.post(f"/api/documents/{document_id}/approve-unflagged").raise_for_status()
        refreshed = client.get(f"/api/documents/{document_id}/segments")
        refreshed.raise_for_status()
        for segment in refreshed.json():
            if segment["status"] != "approved":
                text = segment["translation"]["candidate_text"]
                response = client.post(
                    f"/api/segments/{segment['id']}/approve",
                    json={"text": text, "actor": "demo-validator"},
                )
                response.raise_for_status()

    export = client.post(f"/api/documents/{document_id}/export", json={"format": "docx"})
    export.raise_for_status()
    export_path = Path(export.json()["export_file_uri"])
    feedback_pack = client.get(f"/api/documents/{document_id}/feedback-pack")
    feedback_pack.raise_for_status()
    feedback_manifest = _read_feedback_manifest(feedback_pack.content)
    exported_text = _read_docx_text(export_path)
    ordered_text = [segment.source_text for segment in parse_docx(export_path)]
    return {
        "document_id": document_id,
        "segments": len(segment_payload),
        "memory_applied": process_summary["counts"]["memory_applied"],
        "export_exists": export_path.exists(),
        "export_path": str(export_path),
        "feedback_pack_manifest": feedback_manifest,
        "has_devanagari": _has_devanagari(exported_text),
        "declares_devanagari_font": _docx_declares_font(export_path, DEVANAGARI_FONT),
        "ordered_text_preview": ordered_text[:8],
        "exported_text_preview": exported_text[:12],
    }


def _read_docx_text(path: Path) -> list[str]:
    document = Document(path)
    lines = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                lines.extend(paragraph.text for paragraph in cell.paragraphs if paragraph.text.strip())
    return lines


def _has_devanagari(lines: list[str]) -> bool:
    return any("\u0900" <= char <= "\u097F" for line in lines for char in line)


def _docx_declares_font(path: Path, font_name: str) -> bool:
    with zipfile.ZipFile(path) as archive:
        xml = archive.read("word/document.xml").decode("utf-8")
    return font_name in xml


def _read_feedback_manifest(content: bytes) -> dict:
    with zipfile.ZipFile(io.BytesIO(content)) as archive:
        return json.loads(archive.read("manifest.json").decode("utf-8"))


def _assert_demo_expectations(results: dict[str, dict]) -> None:
    for name, result in results.items():
        assert result["segments"] > 0, name
        assert result["export_exists"], name
        assert result["feedback_pack_manifest"]["approved_segments"] == result["segments"], name
        assert result["has_devanagari"], name
        assert result["declares_devanagari_font"], name
        assert not any("[fixture-" in line for line in result["ordered_text_preview"]), name
    exported_lines = results["public-service-2"]["exported_text_preview"]
    assert "नगरपालिकाले आवेदन ७ दिनभित्र समीक्षा गर्नेछ।" in exported_lines
    assert "वडा नं. ४ ले ठेगाना प्रमाणित गर्नेछ।" in exported_lines
    assert "शुल्क: NPR ५००" in exported_lines
    assert "२०२६-०४-२२" in exported_lines
    assert results["public-service-2"]["memory_applied"] > 0
    assert results["public-service-2"]["feedback_pack_manifest"]["memory_reused_segments"] > 0


if __name__ == "__main__":
    main()
