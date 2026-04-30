from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
import fitz

from sanad_api.services.demo_content import PARAGRAPH_KEYS, TABLE_ROW_KEYS, demo_text, format_date_value, format_phone_value


ROOT = Path(__file__).resolve().parents[3]
DEMO_DIR = ROOT / "samples" / "demo"
FORMAT_CHECKS_DIR = DEMO_DIR / "format-checks"
FONT_NAME = "Noto Sans Devanagari"
PDF_DEVANAGARI_FONT = "/System/Library/Fonts/Supplemental/Devanagari Sangam MN.ttc"


def main() -> None:
    DEMO_DIR.mkdir(parents=True, exist_ok=True)
    FORMAT_CHECKS_DIR.mkdir(parents=True, exist_ok=True)
    create_font_validation(DEMO_DIR / "font-validation.docx")
    create_public_service_fixture(
        DEMO_DIR / "public-service-1.docx",
        language="en",
        applicant="Maya Lama",
        phone="+977-9841234567",
        date="2026-04-21",
        reference="RES-2026-004",
    )
    create_public_service_fixture(
        DEMO_DIR / "public-service-2.docx",
        language="en",
        applicant="Mingmar Tamang",
        phone="+977-9847654321",
        date="2026-04-22",
        reference="RES-2026-005",
    )
    create_public_service_fixture(
        DEMO_DIR / "public-service-cross-scope-proof.docx",
        language="en",
        applicant="Srijana Yonjan",
        phone="+977-9861234567",
        date="2026-04-23",
        reference="RES-2026-006",
    )
    create_public_service_fixture(
        DEMO_DIR / "public-service-nepali-1.docx",
        language="ne",
        applicant="Maya Lama",
        phone="+977-9841234567",
        date="2026-04-21",
        reference="RES-2026-004",
    )
    create_public_service_fixture(
        DEMO_DIR / "public-service-nepali-2.docx",
        language="ne",
        applicant="Mingmar Tamang",
        phone="+977-9847654321",
        date="2026-04-22",
        reference="RES-2026-005",
    )
    create_public_service_fixture(
        DEMO_DIR / "public-service-tamang-1.docx",
        language="tmg",
        applicant="Maya Lama",
        phone="+977-9841234567",
        date="2026-04-21",
        reference="RES-2026-004",
    )
    create_public_service_fixture(
        DEMO_DIR / "public-service-tamang-2.docx",
        language="tmg",
        applicant="Mingmar Tamang",
        phone="+977-9847654321",
        date="2026-04-22",
        reference="RES-2026-005",
    )
    create_tabular_fixture(FORMAT_CHECKS_DIR / "public-service-table-1.csv", date="2026-05-02", reference="RES-2026-004")
    create_tabular_fixture(FORMAT_CHECKS_DIR / "public-service-table-2.csv", date="2026-05-03", reference="RES-2026-005")
    create_tabular_fixture(
        FORMAT_CHECKS_DIR / "public-service-table-1.tsv",
        date="2026-05-02",
        reference="RES-2026-004",
        delimiter="\t",
    )
    create_tabular_fixture(
        FORMAT_CHECKS_DIR / "public-service-table-2.tsv",
        date="2026-05-03",
        reference="RES-2026-005",
        delimiter="\t",
    )
    create_pdf_fixture(FORMAT_CHECKS_DIR / "public-service-notice-1.pdf", date="2026-05-02", reference="RES-2026-004")
    create_pdf_fixture(FORMAT_CHECKS_DIR / "public-service-notice-2.pdf", date="2026-05-03", reference="RES-2026-005")
    print(f"Demo fixtures written to {DEMO_DIR}")


def create_font_validation(path: Path) -> None:
    document = Document()
    document.add_heading("SANAD Font Validation", level=1)
    document.add_paragraph("English: Certificate of Residence Request")
    document.add_paragraph("Nepali: बसोबास प्रमाणपत्र अनुरोध")
    document.add_paragraph("Tamang sample placeholder: तामाङ समुदाय सेवा सूचना")
    document.add_paragraph("Numbers and entities: Ward No. 4, 2026-04-21, +977-9841234567, NPR 500")
    _apply_font(document)
    document.save(path)


def create_public_service_fixture(path: Path, *, language: str, applicant: str, phone: str, date: str, reference: str) -> None:
    document = Document()
    document.add_heading(demo_text("title", language), level=1)
    for key in PARAGRAPH_KEYS:
        document.add_paragraph(demo_text(key, language))

    table = document.add_table(rows=4, cols=2)
    # Use no visible borders — "Table Grid" causes LibreOffice to render
    # black borders during PDF conversion via Gotenberg.
    table.style = "Table Grid"
    _strip_table_borders(table)
    rows = [
        (demo_text(TABLE_ROW_KEYS[0], language), applicant),
        (demo_text(TABLE_ROW_KEYS[1], language), format_phone_value(phone, language) if language != "en" else phone),
        (demo_text(TABLE_ROW_KEYS[2], language), format_date_value(date, language) if language != "en" else date),
        (demo_text(TABLE_ROW_KEYS[3], language), reference),
    ]
    for row, (label, value) in zip(table.rows, rows, strict=True):
        row.cells[0].text = label
        row.cells[1].text = value

    document.add_paragraph(demo_text("fee_500", language))
    _apply_font(document)
    document.save(path)


def create_tabular_fixture(path: Path, *, date: str, reference: str, delimiter: str = ",") -> None:
    lines = [
        delimiter.join(["section", "label", "value"]),
        delimiter.join(["office", "Office", "Residence Support Desk"]),
        delimiter.join(["office", "Date", date]),
        delimiter.join(["office", "Fee", "NPR 500"]),
        delimiter.join(["office", "Reference ID", reference]),
        delimiter.join(["office", "Phone", "+977-01-5550001"]),
        delimiter.join(["office", "Instruction", "Please bring your citizenship card."]),
    ]
    path.write_text("\n".join(lines) + "\n", encoding="utf-8")


def create_pdf_fixture(path: Path, *, date: str, reference: str) -> None:
    document = fitz.open()
    page = document.new_page(width=595, height=842)
    lines = [
        "Residence Certificate Notice",
        "Office: Residence Support Desk",
        f"Date: {date}",
        "Fee: NPR 500",
        f"Reference ID: {reference}",
        "Phone: +977-01-5550001",
        "Please bring your citizenship card.",
    ]
    y = 72
    for index, line in enumerate(lines):
        page.insert_textbox(
            fitz.Rect(56, y, 520, y + 28),
            line,
            fontname="helv",
            fontsize=16 if index == 0 else 12,
            color=(0, 0, 0),
        )
        y += 34 if index == 0 else 26
    # Devanagari sample — only if the host has the font available
    devanagari_font = Path(PDF_DEVANAGARI_FONT)
    if devanagari_font.exists():
        page.insert_font(fontname="sanad-devanagari", fontfile=PDF_DEVANAGARI_FONT)
        page.insert_textbox(
            fitz.Rect(56, y + 12, 520, y + 40),
            "सत्यापन नमुना",
            fontname="sanad-devanagari",
            fontfile=PDF_DEVANAGARI_FONT,
            fontsize=12,
            color=(0, 0, 0),
        )
    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(path)
    document.close()


def _apply_font(document: Document) -> None:
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            _set_run_font(run)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        _set_run_font(run)


def _set_run_font(run) -> None:
    run.font.name = FONT_NAME
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = r_pr._add_rFonts()
    for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
        r_fonts.set(qn(f"w:{attr}"), FONT_NAME)


def _strip_table_borders(table) -> None:
    """Explicitly set all borders to none so LibreOffice won't add defaults."""
    from lxml import etree
    tbl_pr = table._tbl.tblPr
    if tbl_pr is None:
        return
    borders_el = tbl_pr.find(qn("w:tblBorders"))
    if borders_el is not None:
        tbl_pr.remove(borders_el)
    borders = etree.SubElement(tbl_pr, qn("w:tblBorders"))
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = etree.SubElement(borders, qn(f"w:{edge}"))
        border.set(qn("w:val"), "none")
        border.set(qn("w:sz"), "0")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "auto")


if __name__ == "__main__":
    main()
